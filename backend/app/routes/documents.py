import hashlib
import io
import json
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from sqlalchemy.orm import Session
import pypdf
import pdfplumber
import docx
import pdf2image
import pytesseract
from ..database import get_db
from ..models import User, Document, DocumentChunk
from ..schemas import DocumentResponse
from ..auth import get_current_user
from ..services.embedding_service import EmbeddingService

router = APIRouter(prefix="/documents", tags=["Documents Ingestion"])
SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".doc", ".txt")

def format_table_as_markdown(table: list[list[str]]) -> str:
    if not table or not any(table):
        return ""
    lines = []
    # Header
    headers = [str(c or '').strip() for c in table[0]]
    lines.append("| " + " | ".join(headers) + " |")
    lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
    # Rows
    for row in table[1:]:
        row_str = [str(c or '').strip().replace("\n", " ") for c in row]
        if len(row_str) < len(headers):
            row_str += [""] * (len(headers) - len(row_str))
        elif len(row_str) > len(headers):
            row_str = row_str[:len(headers)]
        lines.append("| " + " | ".join(row_str) + " |")
    return "\n".join(lines)

@router.post("/upload", response_model=dict)
async def upload_document(
    file: UploadFile = File(...),
    security_clearance: str = Form("Legal_Officer"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    # Enforcement on Backend: Admin can choose, other roles forced to their own clearance
    if current_user.role != "Admin":
        security_clearance = current_user.role
        
    if security_clearance not in ["Legal_Officer", "Senior_Advisor", "Admin"]:
        raise HTTPException(status_code=400, detail="Invalid clearance level.")

    filename = file.filename or ""
    filename_lower = filename.lower()
    if not filename_lower.endswith(SUPPORTED_EXTENSIONS):
        supported = ", ".join(extension.lstrip(".") for extension in SUPPORTED_EXTENSIONS)
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type. Supported formats: {supported}.",
        )
    
    content_bytes = await file.read()
    file_hash = hashlib.sha256(content_bytes).hexdigest()
    
    existing = db.query(Document).filter(Document.file_hash == file_hash).first()
    if existing:
        raise HTTPException(status_code=400, detail=f"File already uploaded (SHA-256: {file_hash[:10]}).")
    
    doc = Document(
        filename=filename,
        file_hash=file_hash,
        security_clearance=security_clearance,
        uploaded_by=current_user.id
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)
    
    chunks_to_add = []  # List of tuples: (content, is_table, table_json)
    
    try:
        if filename_lower.endswith(".pdf"):
            # Use pdfplumber for high-quality text & table extraction
            with pdfplumber.open(io.BytesIO(content_bytes)) as pdf:
                for page_idx, page in enumerate(pdf.pages):
                    page_text = page.extract_text() or ""
                    
                    # Table extraction
                    tables = page.extract_tables() or []
                    for table in tables:
                        if table and any(table):
                            tbl_md = format_table_as_markdown(table)
                            if tbl_md:
                                chunks_to_add.append((
                                    f"[Table extracted from PDF Page {page_idx+1}]\n{tbl_md}",
                                    True,
                                    table
                                ))
                                
                    if page_text.strip():
                        chunks_to_add.append((
                            f"[PDF Page {page_idx+1}]\n{page_text}",
                            False,
                            None
                        ))
            
            # OCR Fallback: If no content was extracted (e.g., scanned PDF)
            if not chunks_to_add or sum(len(c[0]) for c in chunks_to_add) < 100:
                images = pdf2image.convert_from_bytes(content_bytes)
                for page_idx, img in enumerate(images):
                    ocr_text = pytesseract.image_to_string(img)
                    if ocr_text.strip():
                        chunks_to_add.append((
                            f"[OCR Extracted PDF Page {page_idx+1}]\n{ocr_text}",
                            False,
                            None
                        ))
                        
        elif filename_lower.endswith((".docx", ".doc")):
            doc_obj = docx.Document(io.BytesIO(content_bytes))
            
            # Paragraph text extraction
            paragraphs = [p.text.strip() for p in doc_obj.paragraphs if p.text.strip()]
            full_text = "\n\n".join(paragraphs)
            if full_text:
                chunks_to_add.append((full_text, False, None))
                
            # Tables extraction
            for tbl_idx, table in enumerate(doc_obj.tables):
                grid = []
                for row in table.rows:
                    grid.append([cell.text.strip() for cell in row.cells])
                if grid and any(grid):
                    tbl_md = format_table_as_markdown(grid)
                    if tbl_md:
                        chunks_to_add.append((
                            f"[Table extracted from Word Doc, Table #{tbl_idx+1}]\n{tbl_md}",
                            True,
                            grid
                        ))
                        
        else:
            # Plain text decoding for .txt uploads.
            text_content = content_bytes.decode("utf-8", errors="ignore")
            chunks_to_add.append((text_content, False, None))
            
    except Exception as e:
        db.delete(doc)
        db.commit()
        raise HTTPException(status_code=500, detail=f"Failed to parse document: {str(e)}")
        
    total_text_chunks = 0
    total_table_chunks = 0
    
    for content, is_table, table_json in chunks_to_add:
        sub_chunks = []
        if is_table:
            sub_chunks = [content]  # Tables treated as atomic
        else:
            # Paragraph chunker
            raw_paragraphs = content.split("\n\n")
            current_chunk = ""
            for para in raw_paragraphs:
                para_clean = para.strip()
                if not para_clean:
                    continue
                if len(current_chunk) + len(para_clean) < 1000:
                    current_chunk += "\n\n" + para_clean if current_chunk else para_clean
                else:
                    if current_chunk:
                        sub_chunks.append(current_chunk)
                    current_chunk = para_clean
            if current_chunk:
                sub_chunks.append(current_chunk)
                
        for index, chunk_text in enumerate(sub_chunks):
            vector = await EmbeddingService.get_embedding(chunk_text)
            
            chunk_model = DocumentChunk(
                document_id=doc.id,
                chunk_index=index,
                content=chunk_text,
                is_table=is_table,
                table_json=table_json if is_table else None
            )
            db.add(chunk_model)
            db.commit()
            db.refresh(chunk_model)
            
            # Save vector embedding
            from sqlalchemy import text
            vector_str = "[" + ",".join(map(str, vector)) + "]"
            db.execute(
                text("UPDATE document_chunks SET embedding = :vec WHERE id = :cid"),
                {"vec": vector_str, "cid": chunk_model.id}
            )
            db.commit()
            
            if is_table:
                total_table_chunks += 1
            else:
                total_text_chunks += 1
                
    return {
        "id": doc.id,
        "filename": doc.filename,
        "file_hash": doc.file_hash,
        "security_clearance": doc.security_clearance,
        "uploaded_at": doc.uploaded_at,
        "text_chunks": total_text_chunks,
        "table_chunks": total_table_chunks,
        "ocr_triggered": any("OCR" in c[0] for c in chunks_to_add)
    }

@router.get("/list", response_model=list[DocumentResponse])
def list_documents(db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    from ..auth import ROLE_HIERARCHY
    user_lvl = ROLE_HIERARCHY.get(current_user.role, 1)
    clearances = [r for r, lvl in ROLE_HIERARCHY.items() if lvl <= user_lvl]
    return db.query(Document).filter(Document.security_clearance.in_(clearances)).all()